#!/usr/bin/env python 
# -*- coding:utf-8 -*-
from openpyxl import load_workbook
import datetime
import docx
import re
from docxtpl import DocxTemplate, InlineImage
import win32com.client as client

# 读取刷新后的“阅读数据表”的数据
yesterday = datetime.datetime.now() - datetime.timedelta(days=1)
before_yesterday = datetime.datetime.now() - datetime.timedelta(days=2)
yesterday_str = yesterday.strftime('%m{}%d{}').format("月", "日")
before_yesterday_str = before_yesterday.strftime('%m{}%d{}').format("月", "日")
now = datetime.datetime.now().strftime('%m%d')
in_fppath = 'D:/0刘宇/日报/' + now + '/阅读数据表.xlsx'
out_fppath = 'D:/0刘宇/日报/' + now + '/' + yesterday.strftime('%Y%m%d') + '阅读日报.docx'
w = load_workbook(in_fppath, data_only=True)
wb = w["日"]
cover = wb['J29'].value
trans = wb['J31'].value
quality = wb['J33'].value
h5 = wb['J39'].value
best_new_user = wb['J73'].value
best_opening = wb['J75'].value

# 读取上一期阅读日报运营填写的分析内容
yesterday_word = before_yesterday.strftime('%Y%m%d')
in_word_path = 'D:/0刘宇/日报/' + now + '/' + yesterday_word + '阅读日报.docx'
file = docx.Document(in_word_path)

# 提取上一期的最佳拉新率分析文案
for i in range(len(file.paragraphs)):
    if file.paragraphs[i].text == '【发展转化率最佳】：':
        try:
            best_new_user_kequn = re.search(r'.*?，.*?，(.*)', file.paragraphs[i + 1].text).group(1)
            best_new_user_analysis = file.paragraphs[i + 2].text
        except:
            best_new_user_kequn = ''
            best_new_user_analysis = ''
        break

# 提取上一期的最佳打开率分析文案
for i in range(len(file.paragraphs)):
    if file.paragraphs[i].text == '【打开率最佳】：':
        try:
            best_opening_kequn = re.search(r'.*?，.*?，(.*)', file.paragraphs[i + 1].text).group(1)
            best_opening_analysis = file.paragraphs[i + 2].text
        except:
            best_opening_kequn = ''
            best_opening_analysis = ''
        break

# 提取上一期的明日计划内容
plan = file.paragraphs[-1].text

tpl = DocxTemplate("./日报python模板.docx")

# 将各个内容变量填入Word模板并生成新文件
context = {
    "yesterday": yesterday_str,
    "before_yesterday": before_yesterday_str,
    "cover": cover,
    "trans": trans,
    "quality": quality,
    "h5": h5,
    "best_new_user": best_new_user,
    "best_opening": best_opening,
    "best_new_user_kequn": best_new_user_kequn,
    "best_new_user_analysis": best_new_user_analysis,
    "best_opening_kequn": best_opening_kequn,
    "best_opening_analysis": best_opening_analysis,
    "plan": plan
}

tplrender(context)
tpl.save(out_fppath)
print("Word生成成功").

xl_app = client.gencache.EnsureDispatch("Excel.Application")
xl_app.Visible = False
wb = xl_app.Workbooks.Open(in_fppath)

doc_app = client.gencache.EnsureDispatch("Word.Application")
doc = doc_app.Documents.Open(out_fppath)

rng = doc.Range()
rng.Find.Execute('各活动用户发展效果如下(T-2)：')
rng.MoveStart(Unit=4, Count=1)
wb.Worksheets('日').Shapes(2).Copy()
rng.Paste()

rng = doc.Range()
rng.Find.Execute('各活动打开率如下(T-1)：')
rng.MoveStart(Unit=4, Count=1)
wb.Worksheets('日').Shapes(1).Copy()
rng.Paste()

wb.Close()
doc.Save()
doc.Close()