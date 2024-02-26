#!/usr/bin/env python 
# -*- coding:utf-8 -*-
import datetime
import docx
import re
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
import pandas as pd
import laxin
import dakai
import pymysql

# 资源量
resource = 500000

# 作图*2,并保存至本地
best_laxin_name, best_laxin_rate = laxin.laxin()
best_open_name, best_open_rate = dakai.dakai()

# mysql数据库连接配置
con = pymysql.connect(
    host="10.11.24.182",
    port=3306,
    user='migu',  # 在这里输入用户名
    password='migu123',  # 在这里输入密码
    charset='utf8',
    database='migu'
)

# 读取活动明细.sql文件并生成dateframe
with open('./活动明细.sql', encoding='utf-8') as f:
    hd_sql = f.read()
df_hd = pd.read_sql(hd_sql, con)
df_hd['任务下发时间'] = pd.to_datetime(df_hd['任务下发时间'])
# 读取数据明细.sql文件并生成dateframe
with open('./数据明细.sql', encoding='utf-8') as f:
    sj_sql = f.read()
df_sj = pd.read_sql(sj_sql, con)
df_sj['数据日期'] = pd.to_datetime(df_sj['数据日期'])

# 汇总所有需要用到的日期
now = datetime.datetime.now()
yesterday = datetime.datetime.now() - datetime.timedelta(days=1)
be_yesterday = datetime.datetime.now() - datetime.timedelta(days=2)
be_be_yesterday = datetime.datetime.now() - datetime.timedelta(days=3)
be_be_be_yesterday = datetime.datetime.now() - datetime.timedelta(days=4)

# 指标计算
# 投放活动数
num_act_dis = len(df_hd.loc[df_hd['任务下发时间'] == yesterday.strftime('%Y-%m-%d'), ['三级分类']]['三级分类'].unique())
# 投放活动波次数
num_act_all = len(df_hd.loc[df_hd['任务下发时间'] == yesterday.strftime('%Y-%m-%d'), ['三级分类']]['三级分类'])
# 覆盖次数
num_cover = df_hd.loc[df_hd['任务下发时间'] == yesterday.strftime('%Y-%m-%d'), ['成功发送用户数']]['成功发送用户数'].sum()
# 资源使用率
res_rate = num_cover/resource
# 打开率
open_rate = round((df_hd.loc[df_hd['任务下发时间'] == yesterday.strftime('%Y-%m-%d'), ['页面UV']]['页面UV'].sum())/num_cover, 8)
# 拉新量
num_laxin = df_sj.loc[df_sj['数据日期'] == be_yesterday.strftime('%Y-%m-%d'), ['手机号码拉新量']]['手机号码拉新量'].sum()
# 三日拉新率
rate_laxin_three = round((df_hd.loc[(df_hd['任务下发时间'] >= be_be_be_yesterday.strftime('%Y-%m-%d')) & (df_hd['任务下发时间'] <= be_yesterday.strftime('%Y-%m-%d')), ['拉新量']]['拉新量'].sum())/
                    (df_hd.loc[(df_hd['任务下发时间'] >= be_be_be_yesterday.strftime('%Y-%m-%d')) & (df_hd['任务下发时间'] <= be_yesterday.strftime('%Y-%m-%d')), ['成功发送用户数']]['成功发送用户数'].sum()),8)
# 三日打开率
rate_open_three = round((df_hd.loc[(df_hd['任务下发时间'] >= be_be_yesterday.strftime('%Y-%m-%d')) & (df_hd['任务下发时间'] <= yesterday.strftime('%Y-%m-%d')), ['页面UV']]['页面UV'].sum())/
                    (df_hd.loc[(df_hd['任务下发时间'] >= be_be_yesterday.strftime('%Y-%m-%d')) & (df_hd['任务下发时间'] <= yesterday.strftime('%Y-%m-%d')), ['成功发送用户数']]['成功发送用户数'].sum()),8)
# 次日留存率
nd_retention_rate = round(((df_sj.loc[df_sj['数据日期'] == be_be_yesterday.strftime('%Y-%m-%d'), ['手机号码拉新量']]['手机号码拉新量']*df_sj.loc[df_sj['数据日期'] == be_be_yesterday.strftime('%Y-%m-%d'), ['次日留存率']]['次日留存率']).sum())/
(df_sj.loc[df_sj['数据日期'] == be_be_yesterday.strftime('%Y-%m-%d'), ['手机号码拉新量']]['手机号码拉新量'].sum()),8)
be_retention_rate = round(((df_sj.loc[df_sj['数据日期'] == be_be_be_yesterday.strftime('%Y-%m-%d'), ['手机号码拉新量']]['手机号码拉新量']*df_sj.loc[df_sj['数据日期'] == be_be_be_yesterday.strftime('%Y-%m-%d'), ['次日留存率']]['次日留存率']).sum())/
(df_sj.loc[df_sj['数据日期'] == be_be_be_yesterday.strftime('%Y-%m-%d'), ['手机号码拉新量']]['手机号码拉新量'].sum()),8)
# 次日留存率环比
huanbi = round((nd_retention_rate/be_retention_rate)-1, 8)
# 增长or下降
change_str = '下降' if huanbi < 0 else '增长'

# 读取上一期音乐日报运营填写的分析内容
in_word_path = 'D:/0刘宇/音乐日周报/' + now.strftime('%m%d') + '/' + be_yesterday.strftime('%Y%m%d') + '音乐日报.docx'
file = docx.Document(in_word_path)
out_fppath = 'D:/0刘宇/音乐日周报/' + now.strftime('%m%d') + '/' + yesterday.strftime('%Y%m%d') + '音乐日报.docx'

# 提取上一期的最佳拉新率分析文案
for i in range(len(file.paragraphs)):
    if file.paragraphs[i].text == '【拉新率最佳】：':
        try:
            best_new_user_kequn = re.search(r'.*?，.*?[，；](.*)', file.paragraphs[i + 1].text).group(1)
            best_new_user_analysis = file.paragraphs[i + 2].text
        except:
            best_new_user_kequn = ''
            best_new_user_analysis = ''
        break

# 提取上一期的最佳打开率分析文案
for i in range(len(file.paragraphs)):
    if file.paragraphs[i].text == '【打开率最佳】：':
        try:
            best_opening_kequn = re.search(r'.*?，.*?[，；](.*)', file.paragraphs[i + 1].text).group(1)
            best_opening_analysis = file.paragraphs[i + 2].text
        except:
            best_opening_kequn = ''
            best_opening_analysis = ''
        break

# 提取上一期的明日计划内容
plan = file.paragraphs[-1].text

# 文本格式化生成：
cover = "{}投放{:d}个活动，{:d}波次，覆盖用户{:.2f}万，资源使用率为{:.2%}。".format(yesterday.strftime('%m{}%d{}').format("月", "日"),num_act_dis,num_act_all,num_cover/10000,res_rate)
trans = "{}打开率{:.3%}；{}新增用户{:d}户；三日拉新率{:.3%}，三日打开率{:.3%}。".format(yesterday.strftime('%m{}%d{}').format("月", "日"),open_rate,be_yesterday.strftime('%m{}%d{}').format("月", "日"),int(num_laxin),rate_laxin_three,rate_open_three)
quality = "次日留存率{:.2%}，较上一日{}{:.2%}。".format(nd_retention_rate,change_str,abs(huanbi))
best_new_user = "《{}》，拉新率{:.3%}，".format(best_laxin_name, best_laxin_rate)
best_opening = "《{}》，打开率{:.3%}，".format(best_open_name, best_open_rate)

# word生成
# 使用模板
tpl = DocxTemplate("./日报python模板.docx")

laxinimage = InlineImage(tpl, './拉新.jpg', width=Mm(140))
dakaiimage = InlineImage(tpl, './打开.jpg', width=Mm(140))

# 将各个内容变量填入Word模板并生成新文件
context = {
    "yesterday": yesterday.strftime('%m{}%d{}').format("月", "日"),
    "before_yesterday": be_yesterday.strftime('%m{}%d{}').format("月", "日"),
    "cover": cover,
    "trans": trans,
    "quality": quality,
    "best_new_user": best_new_user,
    "best_opening": best_opening,
    "best_new_user_kequn": best_new_user_kequn,
    "best_new_user_analysis": best_new_user_analysis,
    "best_opening_kequn": best_opening_kequn,
    "best_opening_analysis": best_opening_analysis,
    "plan": plan,
    "laxinimage": laxinimage,
    "dakaiimage": dakaiimage
}

tpl.render(context)
tpl.save(out_fppath)
print("Word生成成功")